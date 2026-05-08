#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import List, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A3, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ACTION_COLORS = {
    "keep": colors.HexColor("#E8F5E9"),
    "hold_for_specialist": colors.HexColor("#FFF8E1"),
    "drop_no_evidence": colors.HexColor("#FCE4EC"),
    "drop_duplicate": colors.HexColor("#F3E5F5"),
    "drop_impossible": colors.HexColor("#FFEBEE"),
}

ACTION_DOCX_COLORS = {
    "keep": "E8F5E9",
    "hold_for_specialist": "FFF8E1",
    "drop_no_evidence": "FCE4EC",
    "drop_duplicate": "F3E5F5",
    "drop_impossible": "FFEBEE",
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def cleanup_inline(text: str) -> str:
    text = text.replace("<br>", "\n")
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def extract_product_key(text: str) -> str:
    match = re.search(r"product_key:\s*([^\n]+)", text)
    return match.group(1).strip() if match else "unknown_product"


def extract_bullets(text: str, heading: str) -> List[str]:
    pattern = rf"## {re.escape(heading)}\n\n((?:- .+\n)+)"
    match = re.search(pattern, text)
    if not match:
        return []
    lines = []
    for raw in match.group(1).splitlines():
        raw = raw.strip()
        if raw.startswith("- "):
            lines.append(cleanup_inline(raw[2:]))
    return lines


def extract_table_rows(text: str) -> List[List[str]]:
    marker = "## Таблица Кандидатов"
    if marker not in text:
        raise ValueError("Не нашла секцию 'Таблица Кандидатов'")
    after = text.split(marker, 1)[1]
    lines = after.splitlines()
    table_lines = []
    started = False
    for line in lines:
        if line.startswith("|"):
            started = True
            table_lines.append(line)
            continue
        if started and not line.strip():
            break
    if len(table_lines) < 3:
        raise ValueError("Не удалось прочитать таблицу кандидатов")

    rows: List[List[str]] = []
    for line in table_lines:
        parts = [cleanup_inline(p) for p in line.strip().strip("|").split("|")]
        rows.append(parts)
    # Skip markdown separator row
    return [rows[0]] + [r for r in rows[2:] if len(r) == len(rows[0])]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.0) -> None:
    cell.text = ""
    for idx, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)


def build_docx(md_path: Path, product_key: str, anchors: Sequence[str], rows: List[List[str]]) -> Path:
    out_path = md_path.with_name(md_path.stem + "__view.docx")
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Техническая Матрица: Быстрый Просмотр")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run(product_key)
    subrun.font.size = Pt(11)

    if anchors:
        anchor_p = doc.add_paragraph()
        anchor_p.add_run("Factual anchors:\n").bold = True
        for line in anchors:
            anchor_p.add_run(f"• {line}\n")
        for run in anchor_p.runs:
            run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False

    widths = [0.55, 1.5, 2.1, 1.8, 1.8, 1.1, 1.15, 2.05]
    for idx, name in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, name, bold=True, size=8.5)
        set_cell_shading(cell, "D9E2F3")
        cell.width = Inches(widths[idx])

    for row in rows[1:]:
        cells = table.add_row().cells
        action = row[6].strip()
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=False, size=7.5)
            if idx == 6:
                set_cell_shading(cells[idx], ACTION_DOCX_COLORS.get(action, "FFFFFF"))
            elif len(table.rows) % 2 == 0:
                set_cell_shading(cells[idx], "FAFAFA")
            cells[idx].width = Inches(widths[idx])

    doc.save(out_path)
    return out_path


def paragraphize(text: str, style: ParagraphStyle) -> Paragraph:
    safe = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return Paragraph(safe, style)


def build_pdf(md_path: Path, product_key: str, anchors: Sequence[str], rows: List[List[str]]) -> Path:
    out_path = md_path.with_name(md_path.stem + "__view.pdf")
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=landscape(A3),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.fontName = "Helvetica-Bold"
    title_style.fontSize = 18
    title_style.leading = 22

    body_style = ParagraphStyle(
        "BodySmall",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=7.5,
        leading=9,
        spaceAfter=0,
    )
    header_style = ParagraphStyle(
        "HeaderSmall",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=8,
        leading=9,
        alignment=1,
    )

    story = [
        Paragraph("Техническая Матрица: Быстрый Просмотр", title_style),
        Spacer(1, 4 * mm),
        Paragraph(product_key, styles["Heading3"]),
    ]

    if anchors:
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph("Factual anchors", styles["Heading4"]))
        for line in anchors:
            story.append(Paragraph(f"• {line}", body_style))

    story.append(Spacer(1, 4 * mm))

    table_data = []
    table_data.append([paragraphize(c, header_style) for c in rows[0]])
    for row in rows[1:]:
        table_data.append([paragraphize(cell, body_style) for cell in row])

    col_widths = [18 * mm, 48 * mm, 63 * mm, 54 * mm, 50 * mm, 28 * mm, 30 * mm, 60 * mm]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BDBDBD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]

    for i, row in enumerate(rows[1:], start=1):
        if i % 2 == 0:
            style_cmds.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#FAFAFA")))
        action = row[6].strip()
        if action in ACTION_COLORS:
            style_cmds.append(("BACKGROUND", (6, i), (6, i), ACTION_COLORS[action]))

    table.setStyle(TableStyle(style_cmds))
    story.append(table)
    doc.build(story)
    return out_path


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: export_matrix_view.py /path/to/matrix.md")
        return 1

    md_path = Path(sys.argv[1]).expanduser().resolve()
    text = read_text(md_path)
    product_key = extract_product_key(text)
    anchors = extract_bullets(text, "Factual Anchors")
    rows = extract_table_rows(text)

    docx_path = build_docx(md_path, product_key, anchors, rows)
    pdf_path = build_pdf(md_path, product_key, anchors, rows)

    print(docx_path)
    print(pdf_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
